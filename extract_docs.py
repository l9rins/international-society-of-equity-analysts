import docx
import os
import glob
import sys

docs_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'documents')
files = sorted(glob.glob(os.path.join(docs_dir, '*.docx')))

output_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'all_content.txt')

with open(output_file, 'w', encoding='utf-8') as f:
    for filepath in files:
        filename = os.path.basename(filepath)
        f.write(f"\n{'='*80}\n")
        f.write(f"FILE: {filename}\n")
        f.write(f"{'='*80}\n")
        
        doc = docx.Document(filepath)
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style = para.style.name if para.style else "Normal"
                f.write(f"[{style}] {text}\n")
        
        for i, table in enumerate(doc.tables):
            f.write(f"\n--- TABLE {i+1} ---\n")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                f.write(" | ".join(cells) + "\n")

print(f"Content extracted to {output_file}")
